from docx import Document
import re
from typing import Dict, Any

def extract_entities_from_docx(uploaded_file) -> Dict[str, Any]:
    doc = Document(uploaded_file)
    
    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) >= 2:
                table_rows.append((cells[0], cells[1]))
            all_text += " ".join(cells) + "\n"
    
    entities = {}
    
    key_mappings = {
        'Party A': 'Party_A',
        'Party B': 'Party_B',
        'Trade Date': 'Trade_Date', 
        'Trade Time': 'Trade_Time',
        'Initial Valuation Date': 'Initial_Valuation_Date',
        'Effective Date': 'Effective_Date',
        'Valuation Date': 'Valuation_Date',
        'Termination Date': 'Termination_Date',
        'Notional Amount': 'Notional_Amount',
        'Notional Amount (N)': 'Notional_Amount',
        'Upfront Payment': 'Upfront_Payment',
        'Coupon': 'Coupon',
        'Coupon (C)': 'Coupon',
        'Barrier': 'Barrier',
        'Barrier (B)': 'Barrier',
        'Underlying': 'Underlying',
        'Exchange': 'Exchange',
        'Business Day': 'Business_Day',
        'Interest Payments': 'Interest_Payments',
        'Initial Price': 'Initial_Price',
        'Initial Price (Shareini)': 'Initial_Price',
        'Sharefinal': 'Share_final',
        'Future Price Valuation': 'Future_Price_Valuation',
        'Calculation Agent': 'Calculation_Agent',
        'ISDA Documentation': 'ISDA_Documentation'
    }
    
    for entity_key in key_mappings.values():
        entities[entity_key] = None
    
    for key, value in table_rows:
        clean_key = re.sub(r'\*+', '', key).strip()
        if clean_key in key_mappings:
            clean_value = re.sub(r'\*+', '', value).strip()
            entities[key_mappings[clean_key]] = clean_value
    
    patterns = {
        'Party_A': r'Party\s+A[:\s]*\*{0,2}\s*([A-Z\s&]+?)\s*(?:\*{0,2}|\n|$)',
        'Party_B': r'Party\s+B[:\s]*\*{0,2}\s*([A-Z\s&]+?)\s*(?:\*{0,2}|\n|$)',
        'Trade_Date': r'Trade\s+Date[:\s]*\*{0,2}\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})',
        'Trade_Time': r'Trade\s+Time[:\s]*\*{0,2}\s*([0-9]{1,2}:[0-9]{2}:[0-9]{2})',
        'Initial_Valuation_Date': r'Initial\s+Valuation\s+Date[:\s]*\*{0,2}\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})',
        'Effective_Date': r'Effective\s+Date[:\s]*\*{0,2}\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})',
        'Valuation_Date': r'(?<!Initial\s)Valuation\s+Date[:\s]*\*{0,2}\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})',
        'Termination_Date': r'Termination\s+Date[:\s]*\*{0,2}\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})',
        'Notional_Amount': r'Notional\s+Amount[^:]*[:\s]*\*{0,2}\s*(EUR\s+[0-9,\.]+\s+million)',
        'Upfront_Payment': r'Upfront\s+Payment[:\s]*\*{0,2}\s*(\*{3}TBD\*{3}[^:]*?on\s+the\s+Effective\s+Date)',
        'Coupon': r'Coupon[^:]*[:\s]*\*{0,2}\s*([0-9\.]+%)',
        'Barrier': r'Barrier[^:]*[:\s]*\*{0,2}\s*([0-9\.]+%)',
        'Underlying': r'Underlying[:\s]*\*{0,2}\s*([^:]*?\([^)]*\))',
        'Exchange': r'Exchange[:\s]*\*{0,2}\s*([A-Z]+)',
        'Business_Day': r'Business\s+Day[:\s]*\*{0,2}\s*([A-Z]+)',
        'Interest_Payments': r'Interest\s+Payments[:\s]*\*{0,2}\s*([A-Za-z]+)\s*(?:\*{0,2}|\n|$)',
        'Initial_Price': r'Initial\s+Price[^:]*[:\s]*\*{0,2}\s*(Official\s+closing[^:]*?Exchange)',
        'Share_final': r'Share[_~]?final[_~]?[:\s]*\*{0,2}\s*(Official\s+closing[^:]*?Exchange)',
        'Future_Price_Valuation': r'Future\s+Price\s+Valuation[:\s]*\*{0,2}\s*([A-Za-z\s]+?)(?:\*{0,2}|\n|$)',
        'Calculation_Agent': r'Calculation\s+Agent[:\s]*\*{0,2}\s*(Party\s+[AB]\s+and\s+Party\s+[AB])',
        'ISDA_Documentation': r'ISDA\s+Documentation[:\s]*\*{0,2}\s*([A-Za-z]+)\s*(?:\*{0,2}|\n|$)'
    }
    
    for field, pattern in patterns.items():
        if not entities.get(field):
            match = re.search(pattern, all_text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip()
                value = re.sub(r'\*+', '', value).strip()
                value = re.sub(r'\s+', ' ', value)
                if value:
                    entities[field] = value
    
    return entities